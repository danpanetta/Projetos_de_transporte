# -*- coding: utf-8 -*-

# ===================================================================================================================================

# TGT410062 - Tarefa 02
# Script para processar geração de viagens
# 2020-09-06

# ===================================================================================================================================

# CONFIGURAÇÕES

# arquivo_dados_atuais = '../Dados_de_entrada/dados_atuais.xlsx'
arquivo_dados_atuais = '/Users/danilopanettadefaria/Documentos/Projects/TGT410062/Tarefa_02/Material/Dados_de_entrada/dados_atuais.xlsx'
arquivo_dados_futuros = '/Users/danilopanettadefaria/Documentos/Projects/TGT410062/Tarefa_02/Material/Dados_de_entrada/dados_futuros.xlsx'
arquivo_saida_docx = '/Users/danilopanettadefaria/Documentos/Projects/TGT410062/Tarefa_02/Material/Resultados/relatorio.docx'
diretorio_saida = '/Users/danilopanettadefaria/Documentos/Projects/TGT410062/Tarefa_02/Material/Resultados'
figura_grafico_prod = '/Users/danilopanettadefaria/Documentos/Projects/TGT410062/Tarefa_02/Material/Resultados/fig_prod.png'
figura_grafico_atr = '/Users/danilopanettadefaria/Documentos/Projects/TGT410062/Tarefa_02/Material/Resultados/fig_atr.png'

# ===================================================================================================================================

import pandas as pd
import time
import statsmodels.formula.api as smf
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
import subprocess
import matplotlib.pyplot as plt
import numpy as np
import scipy


# ===================================================================================================================================

def passa_para_inteiro(input):
    try:
        valor_resultante = round(float(input))

    except:
        # print('Não consegui converter %s' % input)
        valor_resultante = input

    return valor_resultante


def dataframe_to_docxtable(df):
    dimensoes = df.shape

    table = document.add_table(rows=dimensoes[0] + 1, cols=dimensoes[1], style="Table Grid")

    for j in range(dimensoes[1]):
        table.cell(0, j).text = str(passa_para_inteiro(df.columns[j]))

    for i in range(dimensoes[0]):
        for j in range(dimensoes[1]):
            table.cell(i + 1, j).text = str(passa_para_inteiro(str(df.values[i, j])))
            table.cell(i + 1, j).paragraphs[0].paragraph_format.alignment = WD_ALIGN_PARAGRAPH.RIGHT


def executa(input):
    print(input)
    subprocess.run(input, shell=True)


# ===================================================================================================================================

start_time = time.time()

# ===================================================================================================================================

# Objetivo 1

dados_atuais = pd.read_excel(arquivo_dados_atuais)
# print(dados_atuais)

# PRODUCAO

results_prod = smf.ols("producao ~ vol_prod + emprego", data=dados_atuais).fit()
# print(results_prod.summary())
# print(dir(results_prod))
parametros_prod = results_prod.params
# print(parametros_prod)

equacao_producao = "y = %.5f + %.5f * vol_prod + %.5f * emprego" % (
parametros_prod[0], parametros_prod[1], parametros_prod[2])
print(equacao_producao)

r_quadrado_prod = results_prod.rsquared
# print(r_quadrado_prod)
print("r_quadrado = %s" % r_quadrado_prod)

# Atracao

results_atr = smf.ols("atracao ~ populacao + emprego", data=dados_atuais).fit()
# print(results_atr.summary())
# print(dir(results_atr))
parametros_atr = results_atr.params
# print(parametros_atr)

equacao_atracao = "y = %.5f + %.5f * populacao + %.5f * emprego" % (
parametros_atr[0], parametros_atr[1], parametros_atr[2])
print(equacao_atracao)

r_quadrado_atr = results_atr.rsquared
# print(r_quadrado_atr)
print("r_quadrado = %s" % r_quadrado_atr)

# ===================================================================================================================================

# Objetivo 2

dados_futuros = pd.read_excel(arquivo_dados_futuros)

producao_futura = results_prod.predict(dados_futuros)
print(producao_futura)

atracao_futura = results_atr.predict(dados_futuros)
print(atracao_futura)

# ===================================================================================================================================

# Objetivo 3

document = Document()

document.add_heading('Relatório sobre Geração de Viagens', 0)

document.add_heading('Dados de entrada', 1)
dataframe_to_docxtable(dados_atuais)

document.add_heading('Produção', 1)
p = document.add_paragraph(equacao_producao)
p = document.add_paragraph("rˆ2 = %.5f" % r_quadrado_prod)

document.add_heading('Atração', 1)
p = document.add_paragraph(equacao_atracao)
p = document.add_paragraph("rˆ2 = %.5f" % r_quadrado_atr)

document.add_heading('Resultados da projeção', 1)

previsao = pd.DataFrame(columns=['zona', 'producao', 'atracao'])
previsao['zona'] = dados_atuais['zona']
previsao['producao'] = producao_futura
previsao['atracao'] = atracao_futura

dataframe_to_docxtable(previsao)

# document.save(arquivo_saida_docx)

# IMPORTANDO PARA PDF

# executa("/Applications/LibreOffice.app/Contents/MacOS/soffice --headless --convert-to pdf %s --outdir %s " % (arquivo_saida_docx, diretorio_saida))

# ===================================================================================================================================

# Objetivo 4


def f_prod(x, y):
    return parametros_prod[0] + parametros_prod[1] * x + parametros_prod[2] * y

def f_atr(x, y):
    return parametros_atr[0] + parametros_atr[1] * x + parametros_atr[2] * y


fig = plt.figure()
ax = fig.add_subplot(111, projection='3d')

xs = dados_atuais['vol_prod']
ys = dados_atuais['emprego']
zs = dados_atuais['producao']

# ax.scatter(xs, ys, zs, marker='o')

# ax.set_xlabel('vol. de produção de carga')
# ax.set_ylabel('emprego')
# ax.set_zlabel('produção')

# x = np.linspace(1000, 3000)
# y = np.linspace(500, 2000)

z = f_prod(xs, ys)

data = np.c_[xs, ys, z]

# ax.plot(x, y, z, '-')

mn = np.min(data, axis=0)
mx = np.max(data, axis=0)
X,Y = np.meshgrid(np.linspace(mn[0], mx[0], 20), np.linspace(mn[1], mx[1], 20))
XX = X.flatten()
YY = Y.flatten()

# best-fit linear plane (1st-order)
A = np.c_[data[:,0], data[:,1], np.ones(data.shape[0])]
C,_,_,_ = scipy.linalg.lstsq(A, data[:,2])    # coefficients

# evaluate it on grid
Z = C[0]*X + C[1]*Y + C[2]


# plot points and fitted surface using Matplotlib
fig1 =  plt.figure(figsize=(8, 8))
ax = fig1.gca(projection='3d')
ax.plot_surface(X, Y, Z, rstride=1, cstride=1, alpha=0.2)
ax.scatter(data[:,0], data[:,1], data[:,2], c='r', s=50)

ax.set_xlabel('vol. de produção de carga')
ax.set_ylabel('emprego')
ax.set_zlabel('produção')


plt.savefig(figura_grafico_prod, bbox_inches="tight", dpi=200)

plt.clf()

fig = plt.figure()
ax = fig.add_subplot(111, projection='3d')

xs = dados_atuais['populacao']
ys = dados_atuais['emprego']
zs = dados_atuais['atracao']
# ax.scatter(xs, ys, zs, marker='o')

# ax.set_xlabel('população')
# ax.set_ylabel('emprego')
# ax.set_zlabel('atração')

# x = np.linspace(4000, 80000)
# y = np.linspace(500, 2000)

# x,Y = np.meshgrid(np.linspace(mn[0], mx[0], 20), np.linspace(mn[1], mx[1], 20))

# z = parametros_atr[0] + parametros_atr[1] * x + parametros_atr[2] * y

# ax.plot(x, y, z, '-')

z = f_atr(xs, ys)

data = np.c_[xs, ys, z]

# ax.plot(x, y, z, '-')

mn = np.min(data, axis=0)
mx = np.max(data, axis=0)
X,Y = np.meshgrid(np.linspace(mn[0], mx[0], 20), np.linspace(mn[1], mx[1], 20))
XX = X.flatten()
YY = Y.flatten()

# best-fit linear plane (1st-order)
A = np.c_[data[:,0], data[:,1], np.ones(data.shape[0])]
C,_,_,_ = scipy.linalg.lstsq(A, data[:,2])    # coefficients

# evaluate it on grid
Z = C[0]*X + C[1]*Y + C[2]


# plot points and fitted surface using Matplotlib
fig1 =  plt.figure(figsize=(8, 8))
ax = fig1.gca(projection='3d')
ax.plot_surface(X, Y, Z, rstride=1, cstride=1, alpha=0.2)
ax.scatter(data[:,0], data[:,1], data[:,2], c='r', s=50)

ax.set_xlabel('população')
ax.set_ylabel('emprego')
ax.set_zlabel('atração')


plt.savefig(figura_grafico_atr, bbox_inches="tight", dpi=200)

# Carregando os graficos no pdf

document.add_heading('Gráfico do modelo de regressão de produção de viagens', 1)

document.add_picture(figura_grafico_prod)

document.add_heading('Gráfico do modelo de regressão de atração de viagens', 1)
document.add_picture(figura_grafico_atr)

document.add_page_break()

# document.save(arquivo_saida_docx)
# executa("/Applications/LibreOffice.app/Contents/MacOS/soffice --headless --convert-to pdf %s --outdir %s " % (arquivo_saida_docx, diretorio_saida))


# ===================================================================================================================================

# Correcao dos valores de atracao futuro

document.add_heading('Resultados da projeção corrigida', 1)

fator_de_ajuste = producao_futura.sum() / atracao_futura.sum()

correcao = pd.DataFrame(columns=['zona', 'producao', 'atracao_corrigida'])
correcao['zona'] = dados_atuais['zona']
correcao['producao'] = producao_futura
correcao['atracao_corrigida'] = atracao_futura * fator_de_ajuste

p = document.add_paragraph('f = %.3f' % float(fator_de_ajuste))

dataframe_to_docxtable(correcao)

document.save(arquivo_saida_docx)


# ===================================================================================================================================

executa("/Applications/LibreOffice.app/Contents/MacOS/soffice --headless --convert-to pdf %s --outdir %s " % (arquivo_saida_docx, diretorio_saida))

# ===================================================================================================================================


end_time = time.time()
print("Tempo de execução = %s segundos." % (end_time - start_time))
